"""Parseur pour les fichiers Word (.docx et .doc).

Extrait le texte du document Word puis applique les memes regles
d'extraction que le PDF parser (regex pour bulletins, factures, contrats, etc.).

Strategies d'extraction :
1. .docx : python-docx (natif) ou fallback ZIP/XML
2. .doc (ancien format binaire) :
   a) olefile (parsing natif OLE2 sans outil externe)
   b) antiword (rapide, leger)
   c) LibreOffice headless (conversion vers .docx puis extraction)
   Sur OVH Cloud, installer : apt-get install -y antiword
"""

import re
import struct
import logging
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from urssaf_analyzer.parsers.base_parser import BaseParser
from urssaf_analyzer.parsers.pdf_parser import PDFParser
from urssaf_analyzer.models.documents import Document, Declaration

logger = logging.getLogger(__name__)

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import olefile
    HAS_OLEFILE = True
except ImportError:
    HAS_OLEFILE = False

# Detecter les outils disponibles pour le format .doc (ancien Word binaire)
HAS_ANTIWORD = shutil.which("antiword") is not None
HAS_LIBREOFFICE = shutil.which("libreoffice") is not None or shutil.which("soffice") is not None


def _extraire_texte_ole_natif(chemin: Path) -> str:
    """Extraction native du texte d'un fichier .doc via parsing OLE2/WordDocument.

    Lit directement les structures binaires du format Word 97-2003 (OLE2 Compound
    Document) sans outil externe. Extrait le flux WordDocument et decode le texte
    brut (ASCII + Unicode).
    """
    # Methode 1 : olefile (si disponible)
    if HAS_OLEFILE:
        try:
            if not olefile.isOleFile(str(chemin)):
                return ""
            ole = olefile.OleFileIO(str(chemin))
            try:
                # Le texte est dans le stream "WordDocument"
                if ole.exists("WordDocument"):
                    wd = ole.openstream("WordDocument").read()
                    # Extraire le texte brut via le Table Stream
                    # (0Table ou 1Table selon le flag dans le FIB)
                    texte_parts = []

                    # Essayer d'abord les streams de texte courants
                    for stream_name in ["WordDocument", "1Table", "0Table"]:
                        if ole.exists(stream_name):
                            raw = ole.openstream(stream_name).read()
                            # Extraire les sequences de texte lisibles
                            _extraire_texte_brut_stream(raw, texte_parts)

                    if texte_parts:
                        texte = "\n".join(texte_parts)
                        if len(texte.strip()) > 50:
                            logger.info("Extraction .doc via olefile reussie: %s", chemin.name)
                            return texte

            finally:
                ole.close()
        except Exception as e:
            logger.debug("olefile parsing echoue pour %s: %s", chemin.name, e)

    # Methode 2 : parsing binaire direct (sans dependance)
    try:
        with open(chemin, "rb") as f:
            data = f.read()

        # Verifier la signature OLE2 (D0 CF 11 E0 A1 B1 1A E1)
        if data[:8] != b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            return ""

        texte = _extraire_texte_binaire_doc(data)
        if texte and len(texte.strip()) > 50:
            logger.info("Extraction .doc via parsing binaire reussie: %s", chemin.name)
            return texte
    except Exception as e:
        logger.debug("Parsing binaire .doc echoue pour %s: %s", chemin.name, e)

    return ""


def _extraire_texte_brut_stream(raw: bytes, parts: list) -> None:
    """Extrait les sequences de texte lisible d'un flux binaire OLE2."""
    # Extraire le texte ASCII (sequences de >= 4 caracteres imprimables)
    texte_ascii = re.findall(rb"[\x20-\x7e\x0a\x0d\xc0-\xff]{4,}", raw)
    for chunk in texte_ascii:
        try:
            decoded = chunk.decode("cp1252", errors="replace")
            # Filtrer le bruit (sequences de caracteres de controle, metadata XML, etc.)
            if _est_texte_significatif(decoded):
                parts.append(decoded.strip())
        except Exception:
            continue

    # Extraire le texte Unicode (UTF-16LE, courant dans les .doc modernes)
    try:
        texte_utf16 = re.findall(
            rb"(?:[\x20-\x7e\xc0-\xff]\x00){4,}",
            raw,
        )
        for chunk in texte_utf16:
            decoded = chunk.decode("utf-16-le", errors="replace")
            if _est_texte_significatif(decoded):
                parts.append(decoded.strip())
    except Exception:
        pass


def _est_texte_significatif(texte: str) -> bool:
    """Determine si une chaine extraite est du texte significatif (pas du bruit)."""
    if len(texte) < 4:
        return False
    # Ratio de caracteres alphabetiques/espaces vs total
    alpha_space = sum(1 for c in texte if c.isalpha() or c.isspace() or c.isdigit() or c in ".,;:!?'-/()@")
    ratio = alpha_space / len(texte)
    return ratio > 0.6


def _extraire_texte_binaire_doc(data: bytes) -> str:
    """Extraction de texte depuis les donnees brutes d'un fichier .doc OLE2.

    Parse la structure FIB (File Information Block) pour localiser le texte
    dans le flux WordDocument.
    """
    texte_parts = []

    # Methode heuristique : chercher les blocs de texte dans tout le fichier
    # Word stocke le texte en ASCII (cp1252) ou UTF-16LE selon le flag clr.fComplex
    # On extrait les deux types

    # 1. Extraire les longues sequences de texte cp1252
    matches = re.findall(rb"[\x20-\x7e\x0a\x0d\xc0-\xff]{10,}", data)
    for m in matches:
        try:
            decoded = m.decode("cp1252", errors="replace")
            decoded = decoded.replace("\r\n", "\n").replace("\r", "\n")
            if _est_texte_significatif(decoded) and len(decoded) >= 10:
                texte_parts.append(decoded.strip())
        except Exception:
            continue

    # 2. Extraire les longues sequences UTF-16LE
    utf16_matches = re.findall(
        rb"(?:[\x20-\x7e\xc0-\xff]\x00){10,}",
        data,
    )
    for m in utf16_matches:
        try:
            decoded = m.decode("utf-16-le", errors="replace")
            if _est_texte_significatif(decoded) and len(decoded) >= 10:
                texte_parts.append(decoded.strip())
        except Exception:
            continue

    # Dedupliquement et assemblage
    seen = set()
    unique_parts = []
    for p in texte_parts:
        if p not in seen and len(p) > 10:
            seen.add(p)
            unique_parts.append(p)

    return "\n".join(unique_parts)


def _extraire_metadata_doc_ole(chemin: Path) -> dict:
    """Extrait les proprietes du document .doc via OLE2 (auteur, date, titre, etc.)."""
    meta = {}
    if not HAS_OLEFILE:
        return meta
    try:
        if not olefile.isOleFile(str(chemin)):
            return meta
        ole = olefile.OleFileIO(str(chemin))
        try:
            ole_meta = ole.get_metadata()
            if ole_meta.title:
                meta["titre"] = ole_meta.title.decode("cp1252", errors="replace") if isinstance(ole_meta.title, bytes) else str(ole_meta.title)
            if ole_meta.author:
                meta["auteur"] = ole_meta.author.decode("cp1252", errors="replace") if isinstance(ole_meta.author, bytes) else str(ole_meta.author)
            if ole_meta.company:
                meta["societe"] = ole_meta.company.decode("cp1252", errors="replace") if isinstance(ole_meta.company, bytes) else str(ole_meta.company)
            if ole_meta.last_saved_by:
                meta["dernier_auteur"] = ole_meta.last_saved_by.decode("cp1252", errors="replace") if isinstance(ole_meta.last_saved_by, bytes) else str(ole_meta.last_saved_by)
            if ole_meta.create_time:
                meta["date_creation"] = str(ole_meta.create_time)
            if ole_meta.last_saved_time:
                meta["date_modification"] = str(ole_meta.last_saved_time)
            if ole_meta.num_pages:
                meta["nb_pages"] = ole_meta.num_pages
            if ole_meta.num_words:
                meta["nb_mots"] = ole_meta.num_words
        finally:
            ole.close()
    except Exception as e:
        logger.debug("Extraction metadata OLE echouee: %s", e)
    return meta


def _extraire_texte_docx_zip(chemin: Path) -> str:
    """Extraction basique du texte depuis un .docx via ZIP/XML (sans python-docx)."""
    import zipfile
    try:
        with zipfile.ZipFile(chemin, "r") as z:
            if "word/document.xml" not in z.namelist():
                return ""
            xml_content = z.read("word/document.xml").decode("utf-8", errors="replace")
            # Strip XML tags to get plain text
            texte = re.sub(r"<[^>]+>", " ", xml_content)
            texte = re.sub(r"\s+", " ", texte).strip()
            return texte
    except Exception:
        return ""


def _convertir_doc_vers_texte(chemin: Path) -> str:
    """Convertit un fichier .doc (ancien format binaire) en texte.

    Essaie dans l'ordre :
    1. Parsing natif OLE2 (sans outil externe)
    2. antiword (rapide, leger)
    3. LibreOffice headless (conversion vers .docx puis extraction)
    """
    # Methode 1: Parsing natif OLE2 (fonctionne sans outil externe)
    texte = _extraire_texte_ole_natif(chemin)
    if texte and len(texte.strip()) > 50:
        return texte

    # Methode 2: antiword
    if HAS_ANTIWORD:
        try:
            result = subprocess.run(
                ["antiword", "-w", "0", str(chemin)],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                logger.info("Extraction .doc via antiword reussie: %s", chemin.name)
                return result.stdout
        except (subprocess.TimeoutExpired, OSError) as e:
            logger.warning("antiword a echoue pour %s: %s", chemin.name, e)

    # Methode 3: LibreOffice headless -> convertir en .docx
    if HAS_LIBREOFFICE:
        try:
            with tempfile.TemporaryDirectory() as td:
                lo_cmd = "libreoffice" if shutil.which("libreoffice") else "soffice"
                result = subprocess.run(
                    [lo_cmd, "--headless", "--convert-to", "docx",
                     "--outdir", td, str(chemin)],
                    capture_output=True, text=True, timeout=60,
                )
                if result.returncode == 0:
                    converted = Path(td) / (chemin.stem + ".docx")
                    if converted.exists():
                        if HAS_DOCX:
                            doc = docx.Document(str(converted))
                            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                            texte = "\n".join(paragraphs)
                        else:
                            texte = _extraire_texte_docx_zip(converted)
                        if texte.strip():
                            logger.info("Extraction .doc via LibreOffice reussie: %s", chemin.name)
                            return texte
        except (subprocess.TimeoutExpired, OSError) as e:
            logger.warning("LibreOffice a echoue pour %s: %s", chemin.name, e)

    return ""


def _extraire_tableaux_doc(chemin: Path) -> list:
    """Tente d'extraire les tableaux d'un .doc en convertissant d'abord en .docx."""
    if not HAS_LIBREOFFICE or not HAS_DOCX:
        return []
    try:
        with tempfile.TemporaryDirectory() as td:
            lo_cmd = "libreoffice" if shutil.which("libreoffice") else "soffice"
            result = subprocess.run(
                [lo_cmd, "--headless", "--convert-to", "docx",
                 "--outdir", td, str(chemin)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                converted = Path(td) / (chemin.stem + ".docx")
                if converted.exists():
                    doc = docx.Document(str(converted))
                    tableaux = []
                    for table in doc.tables:
                        rows = []
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            rows.append(cells)
                        if rows:
                            tableaux.append(rows)
                    return tableaux
    except Exception:
        pass
    return []


class DocxParser(BaseParser):
    """Parse les fichiers .docx et .doc en extrayant le texte puis en reutilisant le PDF parser.

    Pour les .doc (ancien format binaire), le parsing natif OLE2 est tente en premier
    (sans outil externe), puis antiword, puis LibreOffice. Cela permet d'analyser
    les documents .doc meme sur les serveurs ou antiword/LibreOffice ne sont pas installes.
    """

    def __init__(self):
        self._pdf_parser = PDFParser()

    def peut_traiter(self, chemin: Path) -> bool:
        return chemin.suffix.lower() in (".docx", ".doc")

    def extraire_metadata(self, chemin: Path) -> dict[str, Any]:
        ext = chemin.suffix.lower()
        try:
            texte = self._extraire_texte(chemin)
            meta = {
                "format": "doc" if ext == ".doc" else "docx",
                "nb_caracteres": len(texte),
                "python_docx_disponible": HAS_DOCX,
            }
            if ext == ".doc":
                meta["olefile_disponible"] = HAS_OLEFILE
                meta["antiword_disponible"] = HAS_ANTIWORD
                meta["libreoffice_disponible"] = HAS_LIBREOFFICE
                meta["parsing_natif_ole2"] = True  # Toujours dispo (fallback binaire)
                # Extraire les proprietes OLE (auteur, titre, societe, dates)
                ole_meta = _extraire_metadata_doc_ole(chemin)
                if ole_meta:
                    meta["proprietes_document"] = ole_meta
            if HAS_DOCX and ext == ".docx":
                try:
                    doc = docx.Document(str(chemin))
                    meta["nb_paragraphes"] = len(doc.paragraphs)
                    meta["nb_tableaux"] = len(doc.tables)
                    # Extraire les proprietes du document .docx
                    props = doc.core_properties
                    doc_props = {}
                    if props.title:
                        doc_props["titre"] = props.title
                    if props.author:
                        doc_props["auteur"] = props.author
                    if props.created:
                        doc_props["date_creation"] = str(props.created)
                    if props.modified:
                        doc_props["date_modification"] = str(props.modified)
                    if props.last_modified_by:
                        doc_props["dernier_auteur"] = props.last_modified_by
                    if doc_props:
                        meta["proprietes_document"] = doc_props
                except Exception:
                    pass
            return meta
        except Exception as e:
            return {"format": ext.lstrip("."), "erreur": str(e)}

    def parser(self, chemin: Path, document: Document) -> list[Declaration]:
        texte = self._extraire_texte(chemin)
        tableaux = self._extraire_tableaux(chemin)
        if not texte.strip():
            from urssaf_analyzer.models.documents import Employeur
            decl = Declaration(
                type_declaration="document_non_extractible",
                reference=document.nom_fichier,
                employeur=Employeur(source_document_id=document.id),
                source_document_id=document.id,
                metadata={
                    "type_document": "doc_ancien_format" if chemin.suffix.lower() == ".doc" else "docx",
                    "avertissement": (
                        "Impossible d'extraire le texte de ce document. "
                        "Le fichier est peut-etre protege, corrompu ou ne contient que des images. "
                        "Essayez de le convertir en .docx ou PDF."
                    ),
                },
            )
            return [decl]

        doc_type = self._pdf_parser._detecter_type_document(texte, chemin.name)

        if doc_type == "bulletin":
            return self._pdf_parser._parser_bulletin(texte, tableaux, document)
        elif doc_type == "livre_de_paie":
            return self._pdf_parser._parser_livre_de_paie(texte, tableaux, document)
        elif doc_type == "facture":
            return self._pdf_parser._parser_facture(texte, document)
        elif doc_type == "contrat":
            return self._pdf_parser._parser_contrat(texte, document)
        elif doc_type == "interessement":
            return self._pdf_parser._parser_interessement(texte, document)
        elif doc_type == "attestation":
            return self._pdf_parser._parser_attestation(texte, document)
        else:
            return self._pdf_parser._parser_generique(texte, tableaux, document)

    def _extraire_texte(self, chemin: Path) -> str:
        """Extrait le texte du document Word (.docx ou .doc)."""
        ext = chemin.suffix.lower()

        # Format .doc (ancien binaire) : parsing natif OLE2, antiword ou LibreOffice
        if ext == ".doc":
            return _convertir_doc_vers_texte(chemin)

        # Format .docx
        if HAS_DOCX:
            try:
                doc = docx.Document(str(chemin))
                # Extraire texte des paragraphes + en-tetes/pieds de page
                parts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                # Extraire aussi le texte des en-tetes et pieds de page
                for section in doc.sections:
                    for header_part in [section.header, section.first_page_header, section.even_page_header]:
                        if header_part and header_part.paragraphs:
                            for p in header_part.paragraphs:
                                if p.text.strip():
                                    parts.append(p.text)
                    for footer_part in [section.footer, section.first_page_footer, section.even_page_footer]:
                        if footer_part and footer_part.paragraphs:
                            for p in footer_part.paragraphs:
                                if p.text.strip():
                                    parts.append(p.text)
                return "\n".join(parts)
            except Exception:
                pass
        # Fallback: extraction basique via ZIP
        return _extraire_texte_docx_zip(chemin)

    def _extraire_tableaux(self, chemin: Path) -> list:
        """Extrait les tableaux du document Word au format compatible PDF parser."""
        ext = chemin.suffix.lower()

        # Format .doc : conversion necessaire
        if ext == ".doc":
            return _extraire_tableaux_doc(chemin)

        if not HAS_DOCX:
            return []
        try:
            doc = docx.Document(str(chemin))
            tableaux = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    tableaux.append(rows)
            return tableaux
        except Exception:
            return []
